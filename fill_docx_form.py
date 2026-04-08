import re
from docx import Document
from docx.shared import RGBColor


def extract_docx_fields(docx_path):
    doc = Document(docx_path)
    fields = []

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 2:
                label_text = cells[0].text.strip()
                value_text = cells[1].text.strip()
                if label_text and (value_text == "" or re.match(r"^[_\s]*$", value_text)):
                    fields.append({
                        "field_name": label_text,
                        "location": ("table_cell", table, row, cells[1]),
                        "field_type": "text"
                    })

    for para in doc.paragraphs:
        text = para.text
        match = re.match(r"^(.+?):\s*_{3,}\s*$", text)
        if match:
            label = match.group(1).strip()
            fields.append({
                "field_name": label,
                "location": ("paragraph", para),
                "field_type": "text"
            })

    return fields, doc


def autofill_docx(doc_or_path, fields, match_results, output_path):
    if isinstance(doc_or_path, str):
        fields, doc = extract_docx_fields(doc_or_path)
    else:
        doc = doc_or_path

    match_lookup = {m.get("form_question"): m for m in match_results}

    for f in fields:
        field_name = f["field_name"]
        match = match_lookup.get(field_name)
        if not match:
            continue
        decision = match.get("decision", "").lower()
        if decision not in ["autofill", "review suggested"]:
            continue
        answer = str(match.get("answer", "") or "")
        location = f["location"]

        if location[0] == "table_cell":
            _, table, row, cell = location
            for para in cell.paragraphs:
                for run in para.runs:
                    run.text = ""
            if cell.paragraphs:
                run = cell.paragraphs[0].add_run(answer)
                if decision == "review suggested":
                    run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
            else:
                cell.add_paragraph(answer)

        elif location[0] == "paragraph":
            para = location[1]
            original = para.text
            new_text = re.sub(r"(:\s*)_{3,}\s*$", f"\\1{answer}", original)
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)

    doc.save(output_path)
    print(f"DOCX saved as {output_path}")
